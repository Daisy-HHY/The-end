from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT_PATH = ROOT.parent / "基于EfficientNet-B0迁移学习的四类花卉图像识别系统设计与实现-论文草稿.docx"
CURVE_IMAGE = ROOT / "runs" / "efficientnet_final" / "training_curve.png"
CM_IMAGE = ROOT / "runs" / "efficientnet_final" / "confusion_matrix.png"


def set_cn_font(run, font_name="宋体", size=None, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_chars=True):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    if first_line_chars:
        fmt.first_line_indent = Cm(0.74)


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(text)
    set_cn_font(r, "宋体", 12)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        set_cn_font(r, "黑体", 16, True)
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_cn_font(r, "黑体", 14, True)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        set_cn_font(r, "黑体", 12, True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_center_title(doc: Document, text: str, size: int = 20, bold: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_cn_font(r, "黑体", size, bold)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_cn_font(r, "宋体", 10)
    return p


def build_cover(doc: Document):
    add_center_title(doc, "广 东 财 经 大 学", 22, True)
    add_center_title(doc, "人工智能创新实践课程论文", 18, True)
    doc.add_paragraph()
    add_center_title(doc, "基于EfficientNet-B0迁移学习的四类花卉图像识别系统设计与实现", 18, True)

    info = [
        ("课程名称", "人工智能创新实践"),
        ("论文题目", "基于EfficientNet-B0迁移学习的四类花卉图像识别系统设计与实现"),
        ("班级", "________________"),
        ("姓名", "________________"),
        ("学号", "________________"),
        ("指导教师", "________________"),
        ("完成日期", "2026年6月"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row, (left, right) in zip(table.rows, info):
        row.cells[0].text = left
        row.cells[1].text = right
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_cn_font(run, "宋体", 12)
    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if widths:
                cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_cn_font(run, "宋体", 10.5, row is table.rows[0])
    return table


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    build_cover(doc)

    add_heading(doc, "摘要", 1)
    add_body_paragraph(
        doc,
        "花卉图像识别是计算机视觉领域中具有代表性的细粒度分类任务，在农业辅助识别、植物科普教育、移动端智能识图等场景中具有一定应用价值。原项目采用TensorFlow 1.x框架和浅层卷积神经网络实现四类花卉识别，能够完成基础训练与单张图像预测，但在模型结构、实验规范性和结果表现方面仍有较大的优化空间。为提高系统的识别精度与论文研究价值，本文在保留原项目功能主线的基础上，引入PyTorch框架与EfficientNet-B0迁移学习模型，对原有方案进行改进，并构建了更加规范的训练、验证与测试流程。"
    )
    add_body_paragraph(
        doc,
        "本文首先对四类花卉数据集进行了整理与统计，共包含蒲公英、玫瑰、向日葵和郁金香四个类别3037张图像，并按照2125:608:304的规模划分训练集、验证集和测试集。在模型训练阶段，本文分别构建了随机初始化的EfficientNet-B0基线模型和基于ImageNet预训练权重的迁移学习模型，通过统一的数据预处理、余弦退火学习率调度、标签平滑和分阶段微调策略开展对比实验。实验结果表明，随机初始化基线模型的测试准确率为68.42%，宏平均F1值为66.47%；采用迁移学习后，最终模型测试准确率提升至84.87%，宏平均F1值提升至84.32%，测试准确率相对基线提高24.04%。"
    )
    add_body_paragraph(
        doc,
        "研究结果说明，在中小规模花卉图像数据场景下，基于预训练模型的迁移学习方案能够显著提升分类性能，并增强系统的稳定性与可复现性。本文的工作为课程项目由“能运行”向“有实验支撑、有对比分析、有工程规范”的升级提供了完整实践案例。"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run("关键词：")
    set_cn_font(r1, "黑体", 12, True)
    r2 = p.add_run("花卉图像识别；卷积神经网络；EfficientNet-B0；迁移学习；PyTorch")
    set_cn_font(r2, "宋体", 12)

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景与意义", 2)
    add_body_paragraph(
        doc,
        "随着人工智能技术的发展，图像分类已经从传统手工特征方法逐步过渡到深度学习驱动的方法。与人工设计纹理、边缘或颜色特征相比，卷积神经网络能够通过多层结构自动学习图像中的局部与全局特征，因此在目标识别、医学影像分析、工业视觉检测和植物分类等任务中得到了广泛应用。花卉图像识别作为图像分类中的典型问题，虽然类别数不多，但由于花朵形态相近、背景复杂、拍摄角度多样，仍然具有一定的研究价值。"
    )
    add_body_paragraph(
        doc,
        "从课程实践角度看，花卉识别项目具有数据直观、任务边界清晰、算法路线成熟等优点，适合作为人工智能课程设计的主题。通过该项目，能够系统训练学生在数据整理、模型构建、实验对比、性能分析和应用部署等方面的综合能力。因此，对原有课程项目进行模型升级与实验规范化，不仅可以提高系统识别效果，也有助于形成更完整的课程论文成果。"
    )

    add_heading(doc, "1.2 国内外研究现状", 2)
    add_body_paragraph(
        doc,
        "国外在图像分类领域的研究起步较早。2012年AlexNet在ImageNet竞赛中取得突破后，深度卷积神经网络逐渐成为视觉任务的主流方法。随后，VGG、GoogLeNet、ResNet等网络不断在深度、宽度和残差连接等方面改进模型结构，使图像分类性能持续提升。近年来，模型设计进一步向高效化与轻量化方向发展，EfficientNet通过复合缩放策略在精度和参数量之间取得了较好的平衡。"
    )
    add_body_paragraph(
        doc,
        "国内在植物识别和花卉图像分类方面也开展了较多研究，常见做法包括基于迁移学习对预训练卷积网络进行微调，以及结合数据增强、特征融合和注意力机制提升模型效果。总体而言，对于样本规模较小、类别数量有限的任务，迁移学习能够显著降低训练难度并提升收敛速度，因此已成为课程项目和工程实践中常见且有效的方案。"
    )

    add_heading(doc, "1.3 研究内容与方法", 2)
    add_body_paragraph(
        doc,
        "本文的研究内容主要包括四个方面：第一，梳理原项目的代码结构、数据集组织方式和技术栈，明确其优点与不足；第二，基于PyTorch重新构建图像分类训练流程，引入EfficientNet-B0模型并实现单图预测功能；第三，设计随机初始化与迁移学习两组对比实验，通过统一的数据划分和评估指标验证改进效果；第四，结合实验结果分析模型在不同花卉类别上的识别能力，并总结系统仍可进一步优化的方向。"
    )

    add_heading(doc, "2 相关理论与技术基础", 1)
    add_heading(doc, "2.1 卷积神经网络", 2)
    add_body_paragraph(
        doc,
        "卷积神经网络由卷积层、激活函数、池化层和全连接层等模块组成。卷积层通过局部感受野和权值共享机制提取图像特征，池化层用于降低特征图尺寸并增强平移不变性，全连接层用于综合高层特征完成分类决策。相比传统机器学习依赖人工特征工程的方式，卷积神经网络能够直接从原始图像中学习多层语义表示。"
    )
    add_heading(doc, "2.2 迁移学习", 2)
    add_body_paragraph(
        doc,
        "迁移学习是指将一个任务上已经学习到的知识迁移到另一个相关任务中。在图像分类场景中，常见做法是使用在ImageNet大规模数据集上训练好的模型参数作为初始权重，再在目标数据集上进行微调。由于预训练模型已经学到了较为通用的边缘、纹理和形状表示，因此在小样本任务中通常比从零训练更容易取得较好的效果。"
    )
    add_heading(doc, "2.3 EfficientNet-B0模型", 2)
    add_body_paragraph(
        doc,
        "EfficientNet系列模型的核心思想是通过复合缩放策略同时协调网络深度、宽度和输入分辨率，而不是单独扩大某一个维度。EfficientNet-B0是该系列的基础模型，具有参数量适中、结构较轻、分类性能较优等特点，非常适合中小规模图像分类任务。对于本项目而言，EfficientNet-B0既能体现较新的模型思想，又能够在普通CPU或常规GPU环境下完成训练与验证。"
    )
    add_heading(doc, "2.4 项目实现相关技术", 2)
    add_body_paragraph(
        doc,
        "本文的改进系统采用PyTorch作为深度学习框架，使用torchvision提供的数据变换、预训练模型和图像数据集工具实现训练流程；使用scikit-learn完成数据分层划分和分类指标统计；使用matplotlib输出训练曲线和混淆矩阵图像。与原项目的TensorFlow 1.x静态图方式相比，PyTorch的动态图训练逻辑更直观，便于课程项目后续维护与扩展。"
    )

    add_heading(doc, "3 系统分析与设计", 1)
    add_heading(doc, "3.1 原项目结构分析", 2)
    add_body_paragraph(
        doc,
        "原项目主要由数据读取、浅层CNN模型定义、训练脚本、测试脚本和GUI识别界面组成，整体流程完整，适合课程入门实践。但在深入分析后可以发现，原项目存在模型结构较浅、路径硬编码明显、实验记录不完整、验证集只构建未充分使用等问题，这些问题会直接影响论文中的实验可信度与工程规范性。"
    )
    add_heading(doc, "3.2 数据集构建与预处理", 2)
    add_body_paragraph(
        doc,
        "本项目数据集共包含四类花卉图像3037张，其中蒲公英898张、玫瑰641张、向日葵699张、郁金香799张。考虑到课程项目对复现性的要求，本文在新训练流程中采用固定随机种子进行分层抽样，将数据划分为训练集2125张、验证集608张和测试集304张。训练阶段使用随机裁剪、水平翻转、随机旋转和颜色扰动等数据增强方法，以提高模型的泛化能力；验证与测试阶段则仅进行尺度调整、中心裁剪和标准化处理，以保证结果评估的一致性。"
    )
    add_table(
        doc,
        ["类别", "样本数"],
        [["蒲公英", "898"], ["玫瑰", "641"], ["向日葵", "699"], ["郁金香", "799"], ["总计", "3037"]],
        widths=[2.5, 2.5],
    )
    add_caption(doc, "表3-1 四类花卉数据集样本分布")

    add_heading(doc, "3.3 模型训练流程设计", 2)
    add_body_paragraph(
        doc,
        "为了形成具有说服力的实验结论，本文设计了两组对比实验。第一组为随机初始化的EfficientNet-B0基线模型，用于观察在不使用预训练知识时模型的学习能力；第二组为使用ImageNet预训练权重的迁移学习模型，并在前2轮冻结特征提取层，仅训练分类头部，之后解冻全部网络参数进行联合微调。两组实验均使用相同的数据划分、批大小、输入分辨率和余弦退火学习率调度策略，从而保证对比的公平性。"
    )
    add_heading(doc, "3.4 系统功能模块", 2)
    add_body_paragraph(
        doc,
        "改进后的系统主要包括数据准备模块、模型训练模块、评估分析模块和单图预测模块。数据准备模块负责读取目录式花卉数据并完成分层划分；模型训练模块负责完成EfficientNet-B0的构建、参数优化与权重保存；评估分析模块输出准确率、精确率、召回率、F1值、训练曲线和混淆矩阵；单图预测模块用于读取训练好的模型，对输入花卉图像给出分类概率结果。"
    )

    add_heading(doc, "4 实验设计与结果分析", 1)
    add_heading(doc, "4.1 实验环境", 2)
    add_body_paragraph(
        doc,
        "本文实验在本地计算机的conda环境yolov11_test中完成，Python版本为3.9.23，主要依赖库包括PyTorch、torchvision、scikit-learn、Pillow和matplotlib。由于本地未启用CUDA，本次实验在CPU环境下完成，但这不影响模型对比结果的有效性。"
    )
    add_table(
        doc,
        ["项目", "配置"],
        [
            ["运行环境", "conda环境 yolov11_test"],
            ["Python版本", "3.9.23"],
            ["深度学习框架", "PyTorch + torchvision"],
            ["训练设备", "CPU"],
            ["图像输入尺寸", "160×160"],
            ["批大小", "16"],
        ],
        widths=[2.2, 4.0],
    )
    add_caption(doc, "表4-1 实验环境与主要配置")

    add_heading(doc, "4.2 评价指标", 2)
    add_body_paragraph(
        doc,
        "本文采用准确率（Accuracy）、精确率（Precision）、召回率（Recall）和F1值作为主要评价指标。其中，准确率用于衡量样本总体分类正确的比例；精确率用于衡量预测为某类的样本中实际正确的比例；召回率用于衡量某类真实样本被模型正确识别的比例；F1值则综合反映精确率与召回率的平衡情况。对于多分类任务，宏平均F1值能够更全面地反映模型在各类别上的整体表现。"
    )

    add_heading(doc, "4.3 对比实验方案", 2)
    add_body_paragraph(
        doc,
        "正式实验包括两组方案。其一为随机初始化基线模型，训练8个epoch，不使用预训练权重；其二为迁移学习最终模型，使用ImageNet预训练权重，训练8个epoch，并采用“前2轮冻结、后6轮微调”的策略，同时引入标签平滑与余弦退火学习率调度。两组实验均基于相同训练集、验证集和测试集开展，保证实验结论具有可比性。"
    )
    add_table(
        doc,
        ["实验方案", "权重初始化", "训练策略", "Epoch", "测试准确率"],
        [
            ["基线模型", "随机初始化", "直接训练全部参数", "8", "68.42%"],
            ["最终模型", "ImageNet预训练", "冻结2轮后全量微调", "8", "84.87%"],
        ],
        widths=[1.3, 1.3, 2.2, 0.8, 1.2],
    )
    add_caption(doc, "表4-2 两组实验方案及测试结果对比")

    add_heading(doc, "4.4 实验结果分析", 2)
    add_body_paragraph(
        doc,
        "从实验结果看，随机初始化基线模型在8个epoch后测试准确率达到68.42%，宏平均F1值为66.47%，说明EfficientNet-B0即使在没有预训练权重的情况下也具备一定分类能力。但与最终迁移学习模型相比，基线模型在玫瑰和郁金香类别上的召回率仍然偏低，表明模型在小样本类别和外观相近类别上容易出现混淆。"
    )
    add_body_paragraph(
        doc,
        "采用ImageNet预训练权重后，模型测试准确率提升至84.87%，宏平均F1值提升至84.32%，相较8个epoch的随机初始化基线，测试准确率绝对提高16.45个百分点，相对提升24.04%。该结果说明迁移学习有效利用了大规模数据集上学习到的通用视觉特征，显著提高了模型在本项目中的收敛速度和最终识别精度。"
    )
    add_table(
        doc,
        ["类别", "Precision", "Recall", "F1-score"],
        [
            ["蒲公英", "0.8925", "0.9222", "0.9071"],
            ["玫瑰", "0.7742", "0.7500", "0.7619"],
            ["向日葵", "0.8873", "0.9000", "0.8936"],
            ["郁金香", "0.8205", "0.8000", "0.8101"],
        ],
        widths=[1.6, 1.4, 1.4, 1.4],
    )
    add_caption(doc, "表4-3 最终模型各类别识别指标")

    if CURVE_IMAGE.exists():
        doc.add_picture(str(CURVE_IMAGE), width=Inches(5.8))
        add_caption(doc, "图4-1 最终模型训练与验证曲线")
    if CM_IMAGE.exists():
        doc.add_picture(str(CM_IMAGE), width=Inches(5.2))
        add_caption(doc, "图4-2 最终模型测试集混淆矩阵")

    add_heading(doc, "5 系统总结与可改进方向", 1)
    add_body_paragraph(
        doc,
        "本文完成了原课程项目从传统浅层CNN到EfficientNet-B0迁移学习方案的升级，并实现了数据整理、模型训练、结果分析与单图预测的完整闭环。与原有项目相比，新方案的实验流程更加规范，结果输出更加完整，也更适合作为课程论文写作支撑。"
    )
    add_body_paragraph(
        doc,
        "尽管如此，系统仍有进一步优化空间：第一，可以继续扩充数据集规模，缓解类别不平衡问题；第二，可以尝试更高输入分辨率或更长训练轮数，以进一步挖掘模型性能；第三，可以补充Grad-CAM等可视化解释方法，增强模型决策过程的可解释性；第四，可以在后续工作中引入更轻量的部署方案，将模型与图形界面进一步整合为可直接演示的应用程序。"
    )

    add_heading(doc, "6 结论", 1)
    add_body_paragraph(
        doc,
        "本文围绕四类花卉图像识别任务，对原有课程项目进行了系统梳理与模型升级。研究首先分析了原始TensorFlow 1.x浅层CNN方案在实验设计和识别性能方面的不足，然后基于PyTorch框架构建了EfficientNet-B0迁移学习模型，并通过固定随机种子、分层划分、数据增强、训练曲线和混淆矩阵等手段完善了实验过程。"
    )
    add_body_paragraph(
        doc,
        "实验结果表明，在相同数据集和相近训练预算下，迁移学习模型在测试集上取得84.87%的准确率，明显优于随机初始化基线模型的68.42%，证明预训练特征迁移对于小规模花卉图像分类任务具有显著优势。本文的研究不仅完成了一个较完整的人工智能课程项目，也说明在课程设计中引入较前沿但可落地的模型，可以有效提升项目质量与论文研究深度。"
    )

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] Krizhevsky A, Sutskever I, Hinton G E. ImageNet classification with deep convolutional neural networks[C]//Advances in Neural Information Processing Systems. 2012.",
        "[2] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016.",
        "[3] Tan M, Le Q. EfficientNet: Rethinking model scaling for convolutional neural networks[C]//Proceedings of the 36th International Conference on Machine Learning. 2019.",
        "[4] Deng J, Dong W, Socher R, et al. ImageNet: A large-scale hierarchical image database[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2009.",
        "[5] Abadi M, Agarwal A, Barham P, et al. TensorFlow: Large-scale machine learning on heterogeneous systems[EB/OL]. 2016.",
        "[6] Paszke A, Gross S, Massa F, et al. PyTorch: An imperative style, high-performance deep learning library[C]//Advances in Neural Information Processing Systems. 2019.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, first_line_chars=False)
        r = p.add_run(ref)
        set_cn_font(r, "宋体", 10.5)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
